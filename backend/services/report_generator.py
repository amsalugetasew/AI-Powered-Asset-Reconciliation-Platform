"""
Report generation service for creating PDF, Excel, and Word documents
with charts, analysis, and recommendations.
"""

import io
from datetime import datetime
from typing import Dict, Any, List, Optional
import logging
import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generates reports in multiple formats"""
    
    def __init__(self):
        """Initialize report generator"""
        self.timestamp = datetime.now()
        self.version = "1.0"
    
    def generate_excel_report(
        self,
        filename: str,
        title: str,
        summary_data: Dict[str, Any],
        analysis_results: Dict[str, str],
        chart_data: Optional[Dict[str, Any]] = None,
        records: Optional[List[Dict[str, Any]]] = None
    ) -> io.BytesIO:
        """
        Generate an Excel report
        
        Args:
            filename: Output filename
            title: Report title
            summary_data: Summary statistics
            analysis_results: Analysis content by type
            chart_data: Optional chart data for visualization
            records: Optional detailed records
        
        Returns:
            BytesIO object containing Excel file
        """
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Report"
            
            # Define styles
            title_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            title_font = Font(bold=True, size=16, color="FFFFFF")
            header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            header_font = Font(bold=True, size=11)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Title
            ws.merge_cells('A1:D1')
            title_cell = ws['A1']
            title_cell.value = title
            title_cell.font = title_font
            title_cell.fill = title_fill
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 30
            
            # Report info
            row = 3
            ws[f'A{row}'] = "Report Generated:"
            ws[f'B{row}'] = self.timestamp.strftime("%Y-%m-%d %H:%M:%S")
            row += 1
            
            # Summary section
            ws[f'A{row}'] = "SUMMARY STATISTICS"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            ws[f'A{row}'].fill = header_fill
            row += 1
            
            for key, value in summary_data.items():
                ws[f'A{row}'] = str(key).replace('_', ' ').title()
                ws[f'B{row}'] = value
                ws[f'A{row}'].font = Font(bold=True)
                row += 1
            
            row += 2
            
            # Analysis section
            ws[f'A{row}'] = "ANALYSIS & INSIGHTS"
            ws[f'A{row}'].font = Font(bold=True, size=12)
            ws[f'A{row}'].fill = header_fill
            row += 1
            
            for analysis_type, content in analysis_results.items():
                ws[f'A{row}'] = f"{analysis_type.replace('_', ' ').title()}:"
                ws[f'A{row}'].font = Font(bold=True, italic=True)
                row += 1
                
                # Wrap text for analysis content - ensure content is a string
                content_str = str(content) if content is not None else ""
                # Set value on first cell, then merge the range for proper openpyxl usage
                ws[f'A{row}'].value = content_str
                ws.merge_cells(f'A{row}:D{row}')
                ws[f'A{row}'].alignment = Alignment(wrap_text=True, vertical='top')
                ws.row_dimensions[row].height = None  # Auto height
                row += 2
            
            # Records section
            if records:
                row += 1
                ws[f'A{row}'] = "DETAILED RECORDS (Sample - First 20)"
                ws[f'A{row}'].font = Font(bold=True, size=12)
                ws[f'A{row}'].fill = header_fill
                row += 1
                
                # Create DataFrame and add to Excel
                df = pd.DataFrame(records[:20])
                for r_idx, row_data in enumerate(df.values, start=row + 1):
                    for c_idx, value in enumerate(row_data, start=1):
                        ws.cell(row=r_idx, column=c_idx, value=value)
                
                # Add headers
                for c_idx, col_name in enumerate(df.columns, start=1):
                    cell = ws.cell(row=row, column=c_idx)
                    cell.value = col_name
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.border = border
            
            # Adjust column widths
            for col in ws.columns:
                max_length = 0
                column = get_column_letter(col[0].column)
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                ws.column_dimensions[column].width = min(max_length + 2, 50)
            
            # Save to BytesIO
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error generating Excel report: {str(e)}")
            raise
    
    def generate_pdf_report(
        self,
        filename: str,
        title: str,
        summary_data: Dict[str, Any],
        analysis_results: Dict[str, str],
        chart_data: Optional[Dict[str, Any]] = None,
        records: Optional[List[Dict[str, Any]]] = None
    ) -> io.BytesIO:
        """
        Generate a PDF report
        
        Args:
            filename: Output filename
            title: Report title
            summary_data: Summary statistics
            analysis_results: Analysis content by type
            chart_data: Optional chart data
            records: Optional detailed records
        
        Returns:
            BytesIO object containing PDF file
        """
        if not HAS_REPORTLAB:
            raise ImportError("reportlab is required for PDF generation. Install with: pip install reportlab")
        
        try:
            output = io.BytesIO()
            doc = SimpleDocTemplate(output, pagesize=letter, title=title)
            elements = []
            
            # Custom styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#366092'),
                spaceAfter=12,
                alignment=1  # Center
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#366092'),
                spaceAfter=6,
                spaceBefore=6
            )
            
            # Title
            elements.append(Paragraph(title, title_style))
            elements.append(Spacer(1, 0.3 * inch))
            
            # Report info
            report_info = f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}"
            elements.append(Paragraph(report_info, styles['Normal']))
            elements.append(Spacer(1, 0.2 * inch))
            
            # Summary section
            elements.append(Paragraph("SUMMARY STATISTICS", heading_style))
            summary_data_table = [
                [f"{k.replace('_', ' ').title()}", str(v)]
                for k, v in summary_data.items()
            ]
            summary_table = Table(summary_data_table, colWidths=[3*inch, 3*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#D9E1F2')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F0F0F0')]),
            ]))
            elements.append(summary_table)
            elements.append(Spacer(1, 0.3 * inch))
            
            # Analysis section
            elements.append(Paragraph("ANALYSIS & INSIGHTS", heading_style))
            for analysis_type, content in analysis_results.items():
                elements.append(Paragraph(
                    f"<b>{analysis_type.replace('_', ' ').title()}:</b>",
                    styles['Normal']
                ))
                # Ensure content is a string and handle special characters
                content_str = str(content) if content is not None else ""
                # Escape special XML characters for ReportLab
                para_content = content_str.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br/>')
                elements.append(Paragraph(para_content, styles['Normal']))
                elements.append(Spacer(1, 0.15 * inch))
            
            # Build PDF
            doc.build(elements)
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {str(e)}")
            raise
    
    def generate_word_report(
        self,
        filename: str,
        title: str,
        summary_data: Dict[str, Any],
        analysis_results: Dict[str, str],
        chart_data: Optional[Dict[str, Any]] = None,
        records: Optional[List[Dict[str, Any]]] = None
    ) -> io.BytesIO:
        """
        Generate a Word document report
        
        Args:
            filename: Output filename
            title: Report title
            summary_data: Summary statistics
            analysis_results: Analysis content by type
            chart_data: Optional chart data
            records: Optional detailed records
        
        Returns:
            BytesIO object containing Word file
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx is required for Word generation. Install with: pip install python-docx")
        
        try:
            doc = Document()
            
            # Title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Report info
            doc.add_paragraph(f"Generated: {self.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Summary section
            doc.add_heading('SUMMARY STATISTICS', level=1)
            summary_table = doc.add_table(rows=1, cols=2)
            summary_table.style = 'Light Grid Accent 1'
            
            hdr_cells = summary_table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            for key, value in summary_data.items():
                row_cells = summary_table.add_row().cells
                row_cells[0].text = str(key).replace('_', ' ').title()
                row_cells[1].text = str(value)
            
            # Analysis section
            doc.add_heading('ANALYSIS & INSIGHTS', level=1)
            for analysis_type, content in analysis_results.items():
                doc.add_heading(analysis_type.replace('_', ' ').title(), level=2)
                # Ensure content is a string
                content_str = str(content) if content is not None else ""
                doc.add_paragraph(content_str)
            
            # Records section
            if records:
                doc.add_page_break()
                doc.add_heading('DETAILED RECORDS (Sample)', level=1)
                df = pd.DataFrame(records[:20])
                
                records_table = doc.add_table(rows=1, cols=len(df.columns))
                records_table.style = 'Light Grid Accent 1'
                
                # Headers
                hdr_cells = records_table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                
                # Data
                for _, row_data in df.iterrows():
                    row_cells = records_table.add_row().cells
                    for i, value in enumerate(row_data):
                        row_cells[i].text = str(value)
            
            # Save to BytesIO
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output
            
        except Exception as e:
            logger.error(f"Error generating Word report: {str(e)}")
            raise
    
    def generate_report(
        self,
        format_type: str,
        filename: str,
        title: str,
        summary_data: Dict[str, Any],
        analysis_results: Dict[str, str],
        chart_data: Optional[Dict[str, Any]] = None,
        records: Optional[List[Dict[str, Any]]] = None
    ) -> io.BytesIO:
        """
        Generate report in specified format
        
        Args:
            format_type: 'excel', 'pdf', or 'word'
            filename: Output filename
            title: Report title
            summary_data: Summary statistics
            analysis_results: Analysis content
            chart_data: Optional chart data
            records: Optional detailed records
        
        Returns:
            BytesIO object with report content
        """
        # Validate inputs
        if not isinstance(title, str):
            raise ValueError(f"title must be a string, got {type(title).__name__}")
        if not isinstance(summary_data, dict):
            raise ValueError(f"summary_data must be a dict, got {type(summary_data).__name__}")
        if not isinstance(analysis_results, dict):
            raise ValueError(f"analysis_results must be a dict, got {type(analysis_results).__name__}")
        
        # Convert analysis_results values to strings if needed
        analysis_results = {
            k: str(v) if not isinstance(v, str) else v 
            for k, v in analysis_results.items()
        }
        
        format_type = format_type.lower()
        
        if format_type == 'excel':
            return self.generate_excel_report(filename, title, summary_data, analysis_results, chart_data, records)
        elif format_type == 'pdf':
            return self.generate_pdf_report(filename, title, summary_data, analysis_results, chart_data, records)
        elif format_type == 'word':
            return self.generate_word_report(filename, title, summary_data, analysis_results, chart_data, records)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
